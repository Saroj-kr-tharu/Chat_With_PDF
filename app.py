from dotenv import load_dotenv
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.chat_message_histories import SQLChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain.schema import Document
import faiss
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
import pymupdf
import time
import os
import uuid
from operator import itemgetter
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain.schema.runnable import RunnableMap

class DocumentChat:
    
    def __init__(self):
        """Initialize the DocumentChat class with required configurations."""
        # Load environment variables first
        load_dotenv()
        
        # Environment variables
        self.baseUrl = os.getenv('BASE_URL')
        self.model = os.getenv('MODEL')
        self.temperature = float(os.getenv('Temperature', 0.7))
        self.num_predict = int(os.getenv('Num_predict', 100))
        self.google_api_key = os.getenv('GOOGLE_API_KEY')
        self.db_name = 'health_supplements_VECTOR_DB'
        self.db_history = 'sqlite:///chat_history.db'

        
        # Initialize embeddings
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/text-embedding-004", 
            google_api_key=self.google_api_key, 
            temperature=0.1
        )

    def get_message_history(self, session_id):
        """Retrieve message history for a session.
        
        Args:
            session_id: User or session identifier
            
        Returns:
            List of messages in the chat history
        """
        history = self.get_session_history(session_id)
        return history.messages

    def get_session_history(self, session_id):
        """Get SQL-based chat history for a specific session."""
        return SQLChatMessageHistory(session_id=session_id, connection=self.db_history)
    
    def generate_session_id(self):
        """Generate a random UUID for session identification."""
        return str(uuid.uuid4())

    def is_greeting(self, query):
        """Check if the user input is a greeting"""
        GREETING_PHRASES = ["hi", "hello", "hey", "greetings", "what's up", "howdy"]
        return query.lower().strip() in GREETING_PHRASES

    def format_docs(self, docs):
        """Format documents for context."""
        return "\n\n".join(doc.page_content for doc in docs)

    def create_vector_store(self):
        """Create a new FAISS vector store instance."""
        try:
            # Vector dimension initialization
            index = faiss.IndexFlatL2(len(self.embeddings.embed_query("hello world")))
            vector_store = FAISS(
                embedding_function=self.embeddings,
                index=index,
                docstore=InMemoryDocstore(),
                index_to_docstore_id={}
            )
            return vector_store
        except Exception as e:
            print(f"Error creating vector store: {e}")
            return None

    def vectorStore_create(self, data, DB):
        """Create and save vector store with documents."""
        try:
            vector_store = self.create_vector_store()
            if vector_store is None:
                return None
                
            vector_store.add_documents(documents=data)
            vector_store.save_local(DB)
            return vector_store  # Return the vector store object
        except Exception as e:
            print(f"Error in vectorStore_create: {e}")
            return None

    def vectorStore_search(self, vector_store, question, k=5):
        """Search vector store for similar documents."""
        try:
            response = vector_store.similarity_search(query=question, k=k)
            return response
        except Exception as e:
            print(f"Error in vectorStore_search: {e}")
            return []

    def vectorStore_retrieve(self, vector_store, k=3, fetch_k=20, lambda_mult=1):
        """Get retriever from vector store."""
        try:
            retriever = vector_store.as_retriever(
                search_type='mmr', 
                search_kwargs={'k': k, 'fetch_k': fetch_k, 'lambda_mult': lambda_mult}
            )
            return retriever
        except Exception as e:
            print(f"Error in vectorStore_retrieve: {e}")
            return None

    def vectorStore_load(self, DB):
        """Load existing vector store from disk."""
        try:
            if not os.path.exists(DB):
                return None
                
            instance = FAISS.load_local(DB, embeddings=self.embeddings, allow_dangerous_deserialization=True)
            return instance
        except Exception as e:
            print(f"Error in vectorStore_load: {e}")
            return None


    def ask_gemini(self, question, retriever, session_id):
        """Ask questions using the RAG system."""
        try:
            DEFAULT_GREETING = "I am the informative bot. How can I assist you?"
            if self.is_greeting(question):
                return DEFAULT_GREETING

            llm = ChatGoogleGenerativeAI(
                model="gemini-2.0-flash",
                google_api_key=self.google_api_key,
                temperature=0.1
            )
            
            prompt = ChatPromptTemplate.from_template(
                """You are a helpful and informative bot that answers questions using text from the reference passage included below. 
                Be sure to respond in a complete sentence, being comprehensive, including all relevant background information. 
                However, you are talking to a non-technical audience, so be sure to break down complicated concepts and 
                strike a friendly and conversational tone. 
                If the passage is irrelevant to the answer, you may ignore it.
                QUESTION: '{question}'
                PASSAGE: '{context}'

                ANSWER:"""
            )
            
            # Define a proper function for RunnableMap to use
            def get_context(inputs):
                query = inputs["question"]
                docs = retriever.get_relevant_documents(query)
                return self.format_docs(docs)
            
            # Create the chain using a RunnableMap for proper connection
            chain = (
                RunnableMap({
                    "context": lambda inputs: get_context(inputs),
                    "question": itemgetter("question")
                })
                | prompt
                | llm 
                | StrOutputParser()
            )
            
            # Define function that returns the ChatMessageHistory object
            def get_message_history(session_id):
                return ChatMessageHistory(
                    key=f"chat_history_{session_id}"
                )
            
            # Configure the chain with message history
            chain_with_history = RunnableWithMessageHistory(
                chain,
                lambda session_id: get_message_history(session_id),
                input_messages_key="question",
                history_messages_key="history"
            )
            
            # Get response with history
            response = chain_with_history.invoke(
                {"question": question},
                config={"configurable": {"session_id": session_id}}
            )
            
            return response
            
        except Exception as e:
            print(f"Error in ask_gemini: {e}")
            return f"Sorry, there was an error processing your question: {e}"
    

    def setup_GUI(self):
        """Setup the Streamlit GUI and return retriever if document is uploaded."""
        st.title('Chat With Your Document')

        # Initialize session state for retriever and session_id
        if "retriever" not in st.session_state:
            st.session_state.retriever = None
            
        if "session_id" not in st.session_state:
            st.session_state.session_id = self.generate_session_id()
            st.sidebar.info(f"Session ID: {st.session_state.session_id}")
        
        # Return cached retriever if it exists
        if st.session_state.retriever is not None:
            return st.session_state.retriever

        uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx", "csv"])
        
        if uploaded_file is not None:
            # Check if we already processed this file
            if "uploaded_file_name" in st.session_state and st.session_state.uploaded_file_name == uploaded_file.name:
                return st.session_state.retriever
                
            try:
                # Get file type from file name
                file_type = uploaded_file.name.split('.')[-1].lower()
                
                # Process file based on type
                if file_type == 'pdf':
                    docs = self.process_pdf(uploaded_file)
                elif file_type == 'docx':
                    docs = self.process_docx(uploaded_file)
                elif file_type == 'csv':
                    docs = self.process_csv(uploaded_file)
                else:
                    st.error(f"Unsupported file type: {file_type}")
                    return None
                
                if not docs:
                    st.error(f"No content found in the {file_type.upper()} file.")
                    return None
                
                # Step 2: Split the documents into chunks
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
                chunks = text_splitter.split_documents(docs)
                
                # Step 3: Create vector store
                with st.spinner("Processing document..."):
                    vector_store = self.vectorStore_create(chunks, self.db_name)
                    
                    if vector_store is None:
                        st.error("Failed to create vector store")
                        return None
                
                # Step 4: Create retriever
                retriever = self.vectorStore_retrieve(vector_store)
                if retriever is None:
                    st.error("Failed to create retriever")
                    return None
                
                # Store in session state
                st.session_state.retriever = retriever
                st.session_state.uploaded_file_name = uploaded_file.name
                    
                st.success("Document processed successfully! You can now ask questions.")
            
                return retriever
                
            except Exception as e:
                st.error(f"Error processing document: {e}")
                return None
        
        return None

    def process_pdf(self, uploaded_file):
        """Process PDF file and return documents."""
        bytearray = uploaded_file.read()
        pdf = pymupdf.open(stream=bytearray, filetype="pdf")
        
        docs = []
        
        # Extract text from PDF
        for page in pdf:
            page_text = page.get_text()
            if page_text.strip():  # Only add non-empty pages
                docs.append(Document(page_content=page_text, metadata={"page": page.number}))
        pdf.close()
        
        return docs

    def process_docx(self, uploaded_file):
        """Process DOCX file and return documents."""
        import docx
        
        # Read the content of the uploaded file
        doc = docx.Document(uploaded_file)
        
        # Extract text from paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
        
        docs = []
        if text.strip():
            docs.append(Document(page_content=text, metadata={"source": uploaded_file.name}))
        
        return docs

    def process_csv(self, uploaded_file):
        """Process CSV file and return documents."""
        import pandas as pd
        import io
        
        # Read CSV file
        df = pd.read_csv(io.BytesIO(uploaded_file.getvalue()))
        
        docs = []
        
        # Convert each row to a document
        for index, row in df.iterrows():
            # Convert row to string
            content = "\n".join([f"{col}: {val}" for col, val in row.items()])
            if content.strip():
                docs.append(Document(page_content=content, metadata={"row": index, "source": uploaded_file.name}))
        
        return docs

    def run(self):
        """Main execution function."""
        try:
            # Initialize chat history in session state if it doesn't exist
            if "chat_history" not in st.session_state:
                st.session_state.chat_history = []
            
            # Get or create session ID
            if "session_id" not in st.session_state:
                st.session_state.session_id = self.generate_session_id()
            
            # Show session ID in sidebar (for debugging, can be removed in production)
            # st.sidebar.info(f"Session ID: {st.session_state.session_id}")
            
            retriever = self.setup_GUI()
            
            # Display existing chat history
            for message in st.session_state.chat_history:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])
            
            # Get user input
            prompt = st.chat_input("Type your message here...")
            
            if prompt and retriever is not None:
                # Add user message to chat history
                st.session_state.chat_history.append({"role": "user", "content": prompt})
                
                # Add to SQL chat history
                chat_history = self.get_session_history(st.session_state.session_id)
                chat_history.add_user_message(prompt)
                 
                with st.chat_message("user"):
                    st.markdown(prompt)

                with st.chat_message('assistant'):
                    try:
                        # Create a generator function for streaming
                        def stream_response():
                            response = self.ask_gemini(
                                question=prompt, 
                                retriever=retriever, 
                                session_id=st.session_state.session_id
                            )
                            
                            # Split response by lines to preserve formatting
                            lines = response.split('\n')
                            for line in lines:
                                yield line + '\n'
                                time.sleep(0.1)  # Slightly longer pause between lines

                        # Get the full response
                        response_content = self.ask_gemini(
                            question=prompt, 
                            retriever=retriever, 
                            session_id=st.session_state.session_id
                        )
                            
                        # Stream the response with markdown formatting
                        st.write_stream(stream_response())

                        # Add assistant response to chat history
                        st.session_state.chat_history.append({"role": "assistant", "content": response_content})
                        
                        # Add to SQL chat history
                        chat_history.add_ai_message(response_content)

                    except Exception as e:
                        error_msg = f"Sorry, there was an error processing your question: {e}"
                        st.error(error_msg)
                        st.session_state.chat_history.append({"role": "assistant", "content": error_msg})
            
            elif prompt and retriever is None:
                st.warning("Please upload a PDF document first before asking questions.")
            
        except Exception as e:
            st.error(f"An error occurred during setup: {e}")


# Main entry point
if __name__ == "__main__":
    doc_chat = DocumentChat()
    doc_chat.run()